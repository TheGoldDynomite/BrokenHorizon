from pathlib import Path
from zipfile import ZipFile
from docx import Document

p = Path(r"C:\UnrealProjects\BrokenHorizon\Docs\Broken_Horizon_Game_Design_Document.docx")
d = Document(p)
print("bytes", p.stat().st_size)
print("paragraphs", len(d.paragraphs))
print("tables", len(d.tables))
print("inline_shapes", len(d.inline_shapes))
print("sections", len(d.sections))
print("headings", sum(1 for x in d.paragraphs if x.style.name.startswith("Heading")))
print("page_breaks", sum(x._p.xml.count('w:type="page"') for x in d.paragraphs))
print("empty_table_cells", sum(
    1 for t in d.tables for row in t.rows for c in row.cells if not c.text.strip()
))
with ZipFile(p) as z:
    print("zip_test", z.testzip() or "PASS")
    print("media", len([n for n in z.namelist() if n.startswith("word/media/")]))
