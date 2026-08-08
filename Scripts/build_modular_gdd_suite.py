from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\UnrealProjects\BrokenHorizon")
SOURCE = ROOT / "Docs" / "Broken_Horizon_Game_Design_Document.docx"
OUTDIR = ROOT / "Docs" / "GDD_Modules"
MASTER = OUTDIR / "Broken_Horizon_Modular_GDD_Index.docx"
ZIP_PATH = ROOT / "Docs" / "Broken_Horizon_Modular_GDD_Suite.zip"
HERO = ROOT / "Documentation" / "Art" / "FirstLight_VisualTarget.png"

NAVY = "162735"
STEEL = "35566E"
SLATE = "5E6C76"
ICE = "E8EEF2"
PALE = "F4F6F7"
ORANGE = "C96D2D"
WHITE = "FFFFFF"

MODULES = [
    {
        "heading": "Document Purpose and Design Status",
        "id": "GDD-00",
        "name": "Product Foundation and Design Status",
        "file": "GDD-00_Product_Foundation.docx",
        "owner": "Creative Director / Design Lead",
        "depends": "None",
        "publishes": "Product definition, status language, pillars, target audience, global constraints",
    },
    {
        "heading": "1. Vision, Pillars, and Player Fantasy",
        "id": "GDD-01",
        "name": "Vision, Pillars, and Player Fantasy",
        "file": "GDD-01_Vision_and_Player_Fantasy.docx",
        "owner": "Creative Director",
        "depends": "GDD-00",
        "publishes": "Player promise, pillars, anti-pillars, experience arc",
    },
    {
        "heading": "2. Core Loop and Campaign Structure",
        "id": "GDD-02",
        "name": "Multiplayer Session, Core Loop, and Campaign",
        "file": "GDD-02_Multiplayer_and_Campaign_Loop.docx",
        "owner": "Lead Game Designer / Network Lead",
        "depends": "GDD-00, GDD-01",
        "publishes": "Authority model, session lifecycle, loops, campaign states, failure/respawn policy",
    },
    {
        "heading": "3. Controls, Movement, Interaction, and Combat",
        "id": "GDD-03",
        "name": "Player Controls, Movement, Interaction, and Combat",
        "file": "GDD-03_Player_Combat.docx",
        "owner": "Player Systems Designer",
        "depends": "GDD-01, GDD-02",
        "publishes": "Input actions, movement states, weapon behavior, combat readability",
    },
    {
        "heading": "4. Health, Injuries, Equipment, and Logistics",
        "id": "GDD-04",
        "name": "Health, Injuries, Equipment, and Field Logistics",
        "file": "GDD-04_Health_Equipment_Logistics.docx",
        "owner": "Combat / Economy Designer",
        "depends": "GDD-02, GDD-03",
        "publishes": "Damage states, treatment, field load, supplies, transport economy",
    },
    {
        "heading": "5. AI, Factions, Morale, and Tactical Command",
        "id": "GDD-05",
        "name": "AI, Factions, Morale, and Tactical Command",
        "file": "GDD-05_AI_and_Tactical_Command.docx",
        "owner": "AI Designer / AI Engineer",
        "depends": "GDD-02, GDD-03, GDD-04",
        "publishes": "AI states, perception, morale, field-squad command contract",
    },
    {
        "heading": "6. Missions and Operation First Light",
        "id": "GDD-06",
        "name": "Mission Framework and Operation First Light",
        "file": "GDD-06_Missions_and_First_Light.docx",
        "owner": "Mission Design Lead",
        "depends": "GDD-02 through GDD-05",
        "publishes": "Mission taxonomy, operation rules, First Light sequence and exit criteria",
    },
    {
        "heading": "7. Persistent War, Sectors, Routes, and Strategic Map",
        "id": "GDD-07",
        "name": "Persistent War, Sectors, Routes, and Strategic Map",
        "file": "GDD-07_Persistent_War.docx",
        "owner": "Systems Design Lead",
        "depends": "GDD-02, GDD-04, GDD-06",
        "publishes": "Strategic entities, priority generation, replication/persistence contracts",
    },
    {
        "heading": "8. World, Level, Art, and Audio Direction",
        "id": "GDD-08",
        "name": "World, Level, Art, and Audio Direction",
        "file": "GDD-08_World_Art_Audio.docx",
        "owner": "World Director / Art Director / Audio Director",
        "depends": "GDD-01, GDD-03, GDD-05, GDD-06",
        "publishes": "World grammar, encounter geometry, visual language, audio and performance rules",
    },
    {
        "heading": "9. UI/UX, Onboarding, Accessibility, and Settings",
        "id": "GDD-09",
        "name": "UI/UX, Onboarding, Accessibility, and Settings",
        "file": "GDD-09_UI_UX_Accessibility.docx",
        "owner": "UX Lead",
        "depends": "GDD-02 through GDD-08",
        "publishes": "HUD priority, screen flows, onboarding, multiplayer menu and accessibility requirements",
    },
    {
        "heading": "10. Progression, Difficulty, Economy, and Balance",
        "id": "GDD-10",
        "name": "Progression, Difficulty, Economy, and Balance",
        "file": "GDD-10_Progression_and_Balance.docx",
        "owner": "Systems / Economy Designer",
        "depends": "GDD-02 through GDD-07, GDD-09",
        "publishes": "Progression boundaries, difficulty axes, scoring, balance targets",
    },
    {
        "heading": "11. Technical Architecture, Content Pipeline, Analytics, and QA",
        "id": "GDD-11",
        "name": "Technical Architecture, Content Pipeline, Analytics, and QA",
        "file": "GDD-11_Technical_and_QA.docx",
        "owner": "Technical Director / QA Lead",
        "depends": "All design modules",
        "publishes": "System ownership, network architecture, authoring contracts, telemetry, validation matrix",
    },
    {
        "heading": "12. Production Scope, Roadmap, Risks, and Acceptance Criteria",
        "id": "GDD-12",
        "name": "Production Scope, Roadmap, Risks, and Acceptance",
        "file": "GDD-12_Production_Roadmap.docx",
        "owner": "Producer / Discipline Leads",
        "depends": "GDD-00 through GDD-11",
        "publishes": "Scope tiers, gates, risks, definition of done, open decisions",
    },
    {
        "heading": "Appendix A. Current Baseline Tuning Reference",
        "id": "GDD-13",
        "name": "Canonical Tuning and Content Contracts",
        "file": "GDD-13_Canonical_Contracts.docx",
        "owner": "Design Lead / Technical Design",
        "depends": "GDD-03 through GDD-07, GDD-11",
        "publishes": "Baseline values, First Light IDs and paths, source basis",
    },
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_furniture(doc, label):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    hp = section.header.paragraphs[0]
    hp.clear()
    r = hp.add_run(f"BROKEN HORIZON  /  MODULAR GDD  /  {label}")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    fp = section.footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("VERSION 1.1  |  29 JULY 2026  |  ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    add_page_field(fp)


def add_title_block(doc, module):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(module["id"])
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.add_run(module["name"])
    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    p.add_run("Standalone production module for the cooperative multiplayer game design")
    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    fields = [
        ("Module owner", module["owner"]),
        ("Depends on", module["depends"]),
        ("Publishes", module["publishes"]),
        ("Change rule", "Update this module first, then review every listed consumer in the master index."),
    ]
    for i, (label, value) in enumerate(fields):
        shade(t.cell(i, 0), ICE)
        p0 = t.cell(i, 0).paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(label)
        rr.bold = True
        rr.font.color.rgb = RGBColor.from_string(NAVY)
        p1 = t.cell(i, 1).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.add_run(value)
    doc.add_paragraph()
    h = doc.add_heading("Module Boundary Contract", level=1)
    h.paragraph_format.keep_with_next = True
    for text in [
        "This file may be assigned, reviewed, versioned, or replaced independently.",
        "Terms and IDs published by this module are interfaces; downstream modules consume them without redefining them.",
        "Cross-module changes require an impact review against the master dependency register.",
        "Implementation values remain tunable unless explicitly identified as a persistence, objective, asset-path, or network contract.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Design Specification", level=1)


def is_heading1_paragraph(element):
    if element.tag != qn("w:p"):
        return False
    p_style = element.find("./w:pPr/w:pStyle", namespaces={"w": qn("w:p").split("}")[0].strip("{")})
    return p_style is not None and p_style.get(qn("w:val")) == "Heading1"


def paragraph_text(element):
    return "".join(t.text or "" for t in element.findall(".//w:t", namespaces={"w": qn("w:t").split("}")[0].strip("{")}))


def is_page_break_only(element):
    if element.tag != qn("w:p"):
        return False
    text = paragraph_text(element).strip()
    return not text and 'w:type="page"' in element.xml


def source_blocks(source_doc):
    blocks = {}
    current = None
    for element in source_doc._element.body:
        if element.tag == qn("w:sectPr"):
            continue
        if is_heading1_paragraph(element):
            current = paragraph_text(element)
            blocks[current] = []
            continue
        if current is not None and not is_page_break_only(element):
            blocks[current].append(deepcopy(element))
    return blocks


def append_exit_checklist(doc, module):
    doc.add_heading("Module Exit Checklist", level=1)
    checks = [
        "The module can be read without another file to understand its owned decisions.",
        "Every external dependency is named in the header and not silently duplicated.",
        "Every published interface uses stable terminology consistent with the master index.",
        "Multiplayer authority, replication, reconnect, and persistence implications are addressed where relevant.",
        "Acceptance criteria are testable and assignable to a discipline.",
        "Changes to objective IDs, persistence IDs, reflected names, paths, or save fields include a consumer audit.",
    ]
    for item in checks:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f"End of {module['id']}  /  {module['name']}")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(STEEL)


def build_module(template_doc, blocks, module):
    doc = Document(SOURCE)
    clear_body(doc)
    set_furniture(doc, module["id"])
    add_title_block(doc, module)
    if module["heading"] not in blocks:
        raise KeyError(f"Missing source section: {module['heading']}")
    for element in blocks[module["heading"]]:
        doc._element.body.insert(-1, deepcopy(element))
    append_exit_checklist(doc, module)
    doc.core_properties.title = f"Broken Horizon {module['id']} - {module['name']}"
    doc.core_properties.subject = "Separable modular game design document"
    doc.core_properties.author = "Broken Horizon Development Team"
    path = OUTDIR / module["file"]
    doc.save(path)
    strip_unused_document_images(path)
    return path


def strip_unused_document_images(path):
    """Remove the source template's unreferenced cover image from a module."""
    tmp = path.with_suffix(".optimized.docx")
    rels_name = "word/_rels/document.xml.rels"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    with ZipFile(path, "r") as zin:
        rels_root = ET.fromstring(zin.read(rels_name))
        removed_targets = set()
        for rel in list(rels_root):
            rel_type = rel.attrib.get("Type", "")
            if rel_type.endswith("/image"):
                removed_targets.add("word/" + rel.attrib["Target"].lstrip("/"))
                rels_root.remove(rel)
        rels_xml = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in removed_targets:
                    continue
                data = rels_xml if item.filename == rels_name else zin.read(item.filename)
                zout.writestr(item, data)
    tmp.replace(path)


def build_master():
    doc = Document(SOURCE)
    clear_body(doc)
    set_furniture(doc, "MASTER INDEX")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("MODULAR GAME DESIGN DOCUMENT")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p = doc.add_paragraph(style="Title")
    p.add_run("BROKEN HORIZON\nMODULE INDEX")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Separable specifications for a cooperative multiplayer persistent-war FPS")
    if HERO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(HERO), width=Inches(6.65))
    doc.add_heading("How to Use This Suite", level=1)
    for item in [
        "Start with GDD-00 and GDD-01 for product intent; assign later modules by discipline.",
        "A module owns only the decisions listed in its Publishes field.",
        "Dependencies are read-only inputs. If a dependent rule changes, update the owner module and review all consumers.",
        "GDD-13 centralizes numeric baselines and canonical IDs so other modules reference rather than duplicate contracts.",
        "The original integrated GDD remains the narrative reference; this suite is the production-control form.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Module Register", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Module", "Owner", "Depends on", "File"]
    for i, value in enumerate(headers):
        shade(t.cell(0, i), NAVY)
        r = t.cell(0, i).paragraphs[0].add_run(value)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for idx, module in enumerate(MODULES):
        cells = t.add_row().cells
        values = [module["id"], module["name"], module["owner"], module["depends"], module["file"]]
        for i, value in enumerate(values):
            if idx % 2:
                shade(cells[i], PALE)
            cells[i].paragraphs[0].add_run(value)
    doc.add_heading("Dependency Flow", level=1)
    flows = [
        ("Foundation", "GDD-00 -> GDD-01 -> every downstream module"),
        ("Playable operation", "GDD-02 + GDD-03 + GDD-04 + GDD-05 -> GDD-06"),
        ("Persistent campaign", "GDD-02 + GDD-04 + GDD-06 -> GDD-07"),
        ("Presentation", "GDD-01 + GDD-03 + GDD-05 + GDD-06 -> GDD-08 -> GDD-09"),
        ("Balance", "GDD-02 through GDD-07 + GDD-09 -> GDD-10"),
        ("Implementation proof", "All design modules -> GDD-11"),
        ("Planning and release", "GDD-00 through GDD-11 -> GDD-12"),
        ("Shared contracts", "GDD-13 is consumed by gameplay, mission, war, technical, and QA work"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(["Stream", "Flow"]):
        shade(t.cell(0, i), NAVY)
        r = t.cell(0, i).paragraphs[0].add_run(value)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for idx, (stream, flow) in enumerate(flows):
        cells = t.add_row().cells
        if idx % 2:
            shade(cells[0], PALE)
            shade(cells[1], PALE)
        cells[0].paragraphs[0].add_run(stream)
        cells[1].paragraphs[0].add_run(flow)
    doc.add_heading("Change-Control Rules", level=1)
    for item in [
        "Gameplay tuning changes: update GDD-13, then review the owning gameplay module and GDD-10.",
        "Network authority or session changes: update GDD-02 and GDD-11, then review every interactive module.",
        "Objective or persistence contract changes: update GDD-06/GDD-07 and GDD-13; search all C++, Blueprint, Python, config, and save consumers.",
        "Visual or UX changes: update GDD-08/GDD-09 and require manual editor or playtest review.",
        "Scope changes: update GDD-12 only after affected owner modules accept the new exit criteria.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.core_properties.title = "Broken Horizon - Modular GDD Index"
    doc.core_properties.subject = "Master dependency and ownership index"
    doc.core_properties.author = "Broken Horizon Development Team"
    doc.save(MASTER)


def validate_docx(path):
    with ZipFile(path) as z:
        bad = z.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX {path}: {bad}")
    doc = Document(path)
    empty_cells = sum(
        1 for t in doc.tables for row in t.rows for cell in row.cells
        if not cell.text.strip()
    )
    if empty_cells:
        raise RuntimeError(f"{path.name} has {empty_cells} empty table cells")
    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading")),
    }


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    source_doc = Document(SOURCE)
    blocks = source_blocks(source_doc)
    paths = [build_module(source_doc, blocks, module) for module in MODULES]
    build_master()
    all_paths = [MASTER] + paths
    results = [validate_docx(path) for path in all_paths]
    with ZipFile(ZIP_PATH, "w", ZIP_DEFLATED) as z:
        for path in all_paths:
            z.write(path, arcname=path.name)
    with ZipFile(ZIP_PATH) as z:
        if z.testzip():
            raise RuntimeError("ZIP bundle failed integrity check")
    print(f"master={MASTER}")
    print(f"bundle={ZIP_PATH}")
    for result in results:
        print(
            f"{result['file']} bytes={result['bytes']} "
            f"paragraphs={result['paragraphs']} tables={result['tables']} "
            f"headings={result['headings']}"
        )


if __name__ == "__main__":
    main()
