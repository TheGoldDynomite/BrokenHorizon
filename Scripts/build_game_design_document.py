from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\UnrealProjects\BrokenHorizon")
OUT = ROOT / "Docs" / "Broken_Horizon_Game_Design_Document.docx"
HERO = ROOT / "Documentation" / "Art" / "FirstLight_VisualTarget.png"

NAVY = "162735"
STEEL = "35566E"
SLATE = "5E6C76"
ICE = "E8EEF2"
PALE = "F4F6F7"
ORANGE = "C96D2D"
WHITE = "FFFFFF"
BLACK = "172026"
GREEN = "386641"
RED = "9B2C2C"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ("Title", 32, NAVY, 0, 8),
    ("Subtitle", 15, SLATE, 0, 10),
    ("Heading 1", 19, NAVY, 14, 7),
    ("Heading 2", 14, STEEL, 11, 5),
    ("Heading 3", 11, ORANGE, 8, 3),
]:
    s = styles[name]
    s.font.name = "Aptos Display" if name != "Normal" else "Aptos"
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Number"]:
    s = styles[list_name]
    s.font.name = "Aptos"
    s.font.size = Pt(10)
    s.paragraph_format.left_indent = Inches(0.28)
    s.paragraph_format.first_line_indent = Inches(-0.18)
    s.paragraph_format.space_after = Pt(3)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, NAVY)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(9)
        set_cell_margins(c)
    set_repeat_table_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if ri % 2:
                shade(cells[i], PALE)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(9)
            set_cell_margins(cells[i])
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def callout(label, text, fill=ICE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    c = t.cell(0, 0)
    c.width = Inches(6.75)
    shade(c, fill)
    set_cell_margins(c, 150, 180, 150, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def page():
    doc.add_page_break()

def add_page_num(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)

# Header / footer
header = sec.header
hp = header.paragraphs[0]
hp.text = "BROKEN HORIZON  /  GAME DESIGN DOCUMENT"
hp.style = normal
hp.runs[0].font.size = Pt(8)
hp.runs[0].font.bold = True
hp.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rr = fp.add_run("CONFIDENTIAL - WORKING DESIGN  |  ")
rr.font.size = Pt(8)
rr.font.color.rgb = RGBColor.from_string(SLATE)
add_page_num(fp)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("GAME DESIGN DOCUMENT")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(ORANGE)
p = doc.add_paragraph(style="Title")
p.add_run("BROKEN\nHORIZON")
p.paragraph_format.space_after = Pt(4)
p = doc.add_paragraph(style="Subtitle")
p.add_run("A cooperative near-future tactical FPS shaped by a shared persistent regional war")
if HERO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.add_run().add_picture(str(HERO), width=Inches(6.85))
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Production GDD  |  Version 1.1  |  Multiplayer correction  |  29 July 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(SLATE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Status: implementation-informed design baseline")
r.bold = True
r.font.color.rgb = RGBColor.from_string(STEEL)
page()

h1("Document Purpose and Design Status")
body("This document defines the intended player experience, production boundaries, core systems, content model, and measurable quality targets for Broken Horizon. It is grounded in the current Unreal Engine 5.8 project and its Operation First Light vertical slice. It separates implemented foundations from proposed product-level direction so design, engineering, art, audio, and QA can use one shared baseline.")
callout("Design thesis", "Every firefight is local, but its consequences travel. The player wins by reading terrain, preserving combat effectiveness, and influencing a war that continues beyond a single mission.")
table(["Status", "Meaning", "GDD treatment"], [
    ("Implemented foundation", "Source-backed gameplay or project configuration exists.", "Specified as the current baseline; tuning remains subject to playtest."),
    ("Partially implemented", "The core system exists, but breadth, presentation, or content coverage is incomplete.", "Defined with a production-ready target state."),
    ("Proposed", "A coherent extension of the current architecture, not yet proven in the build.", "Explicitly framed as a target or decision."),
], [1.35, 2.55, 2.85])
h2("Product Snapshot")
table(["Dimension", "Direction"], [
    ("Genre", "Cooperative multiplayer tactical first-person shooter with a shared persistent strategic simulation."),
    ("Setting", "Fictional near-future Eastern European mountain region; cold, wet, politically unbranded, materially believable."),
    ("Platform", "Networked PC-first; controller support, dedicated/listen-server compatibility, and scalable rendering are requirements. Console is a later feasibility gate."),
    ("Audience", "Squads that value coordination, readable lethality, deliberate movement, systemic missions, persistent consequences, and grounded military atmosphere."),
    ("Session", "Squad-based network session; 20-45 minute operation or 60-120 minute extended campaign session. Exact supported player cap is a product decision."),
    ("Business model", "Premium boxed/digital game. No power monetization; no live-service dependency."),
], [1.45, 5.3])
h2("Table of Contents")
for item in [
    "1. Vision, pillars, and player fantasy", "2. Core loop and campaign structure",
    "3. Controls, movement, interaction, and combat", "4. Health, injuries, equipment, and logistics",
    "5. AI, factions, morale, and tactical command", "6. Missions and Operation First Light",
    "7. Persistent war, sectors, routes, and strategic map", "8. World, level, art, and audio direction",
    "9. UI/UX, onboarding, accessibility, and settings", "10. Progression, difficulty, economy, and balance",
    "11. Technical architecture, content pipeline, analytics, and QA", "12. Production scope, roadmap, risks, and acceptance criteria",
]:
    bullet(item)
page()

h1("1. Vision, Pillars, and Player Fantasy")
h2("High Concept")
body("Broken Horizon places a player squad inside a fractured mountain theater where territory, supply routes, militia readiness, and battlefield losses persist. No participant is a super-soldier or omniscient commander. Players are capable field operators whose coordination connects the tactical and strategic layers: securing a relay station can restore intelligence; saving a transport preserves future supply; suffering casualties may win today while weakening tomorrow.")
h2("Player Fantasy")
body("Be part of the field squad that stabilizes a collapsing front through competence rather than spectacle: observe, plan, divide responsibilities, infiltrate, fight, treat teammates, recover what matters, and collectively decide which crisis deserves scarce attention.")
h2("Design Pillars")
table(["Pillar", "Promise", "Design test"], [
    ("Grounded tactical clarity", "Weapons, movement, cover, visibility, and wounds create readable cause and effect.", "Can the player explain why they succeeded or failed without hidden rules?"),
    ("Persistent consequence", "Territory, supply, casualties, objectives, and assets survive travel, death, reconnect, and server travel.", "Does a shared operation outcome materially alter later choices?"),
    ("Systemic operations", "Attack, defense, raid, escort, rescue, and resupply emerge from shared systems.", "Can one ruleset create multiple stories without bespoke scripting?"),
    ("Pressure without chaos", "The war moves, but the player receives legible priorities, approach windows, and recovery opportunities.", "Is urgency meaningful without becoming constant noise?"),
    ("Restrained authenticity", "The world feels plausible and serious without simulating paperwork or fetishizing real factions.", "Does detail serve decisions, atmosphere, or feedback?"),
], [1.4, 2.55, 2.8])
h2("Anti-Pillars")
for x in [
    "No bullet-sponge enemies, randomized gear scores, or rarity colors that override physical logic.",
    "No omnipresent waypoint tunnel; navigation support must preserve observation and terrain reading.",
    "No consequence-free team-wipe restart as the default campaign language.",
    "No neon science-fiction aesthetic, real-world flags, or political faction analogues.",
    "No strategic layer that plays itself while merely reporting flavor text.",
]:
    bullet(x)
h2("Experience Arc")
table(["Phase", "Player emotion", "Primary verbs"], [
    ("Orient", "Uncertainty becoming shared comprehension", "Read map, check status, agree priority, assign roles, prepare"),
    ("Approach", "Controlled tension", "Travel, scout, avoid, position, command"),
    ("Contact", "Focused danger", "Identify, suppress, flank, breach, treat"),
    ("Resolve", "Collective relief with cost awareness", "Regroup, secure, sabotage, rescue, extract, recover"),
    ("Consequences", "Ownership of the campaign", "Review losses, update front, resupply, reprioritize"),
], [1.1, 2.4, 3.25])
page()

h1("2. Core Loop and Campaign Structure")
h2("Multiplayer Session Model")
body("Broken Horizon is a cooperative multiplayer game. The server owns the authoritative war, operation, combat, transport, objective, and persistence state; connected clients receive replicated snapshots and presentation. Players join the same regional campaign, can converge on an active operation, and share its strategic consequences. AI field support complements the human squad but does not redefine the product as a single-player experience.")
table(["Concern", "Design requirement"], [
    ("Authority", "Host or dedicated server validates combat, interactions, transport use, commands, operation outcomes, and save writes."),
    ("Shared war", "All clients see the same replicated sector, supply, priority, and active-operation state revision."),
    ("Operation participation", "Every connected living player can join approach and active phases; objectives evaluate the participant set rather than one designated hero."),
    ("Join in progress", "Joining players receive current war and operation snapshots, then spawn or redeploy at a safe valid field location."),
    ("Disconnect/reconnect", "A disconnect must release claimed seats and interactions safely; reconnect restores the player without duplicating inventory or campaign effects."),
    ("Respawn", "Multiplayer field respawn returns a player to a valid friendly sector or fallback point while preserving the shared operation state."),
    ("Persistence owner", "Only server authority writes or loads the campaign; host loads use server travel so connected clients remain in the session."),
    ("Communication", "Low-friction ping/context commands are required; voice support and platform integration remain product decisions."),
], [1.65, 5.1])
h2("Moment-to-Moment Loop")
for x in [
    "Observe: interpret sightlines, sound, cover, patrols, objectives, and friendly state.",
    "Decide: communicate a route, role, engagement distance, and acceptable risk.",
    "Act: move, cover teammates, lean, crouch or go prone, shoot, throw, interact, treat, and command.",
    "Reassess: share feedback from suppression, injuries, ammunition, casualties, separated players, and enemy behavior.",
]:
    numbered(x)
h2("Operation Loop")
for x in [
    "Receive a strategic priority or discover a local opportunity.",
    "Select an operation by squad agreement and review sector ownership, supply source, travel time, and expected force package.",
    "Prepare ammunition, medical supplies, grenades, roles, AI support, and transport seats.",
    "Travel through the open region; ambient war events can alter the approach.",
    "Execute the operation through coordinated or split tactical routes.",
    "Secure, sabotage, defend, escort, rescue, or extract.",
    "Persist the result, casualties, consumed resources, sector state, and surviving field assets.",
]:
    numbered(x)
h2("Campaign Loop")
body("Campaign time advances in discrete strategic turns triggered by authoritative operation and travel milestones rather than a real-time clock that punishes players for reading shared information. Each turn evaluates contested sectors, supply connectivity, reinforcement capacity, militia pressure, recovery, and new priorities. The squad can influence but not perfectly control every front.")
callout("Pacing rule", "The squad should face one clear primary crisis and at most two credible secondary opportunities. More events may exist in simulation, but the interface must not present all of them as equal emergencies.")
h2("Campaign States")
table(["State", "Condition", "Player-facing effect"], [
    ("Stable front", "Supply routes connected; no critical sector threatened.", "Space to scout, raid, recover, and improve position."),
    ("Contested", "Enemy pressure meets friendly control.", "Attack/defense opportunities, higher patrol density, disrupted travel."),
    ("Breakthrough risk", "A critical connected sector is under-supplied and pressured.", "Short approach window; losing may isolate downstream sectors."),
    ("Encircled", "No friendly supply route reaches the sector.", "Reduced support, scarce resupply, evacuation or relief priorities."),
    ("Regional collapse", "Key command/logistics conditions fail.", "Campaign loss sequence or costly final recovery operation."),
    ("Operational victory", "Player side meets theater objectives.", "Campaign resolution based on territory, force preservation, and civilian impact."),
], [1.15, 2.55, 3.05])
h2("Failure Philosophy")
body("An individual tactical death enters multiplayer field-respawn flow; surviving teammates can continue the operation. A full squad wipe, retreat, or partial success normally preserves the shared campaign outcome instead of rewinding every client. A host-authorized restart may be offered for onboarding or custom sessions, but the distinction must be explicit: restarting rewinds the operation, while accepting the outcome commits the authoritative war state.")
page()

h1("3. Controls, Movement, Interaction, and Combat")
h2("Core Input")
table(["Action", "Keyboard/mouse target", "Controller target", "Notes"], [
    ("Move / look", "WASD / mouse", "Left / right stick", "Full remapping, invert axes, sensitivity per aim state."),
    ("Sprint", "Left Shift", "Left stick click", "Consumes stamina; blocked when medically compromised."),
    ("Crouch / prone", "C / hold C", "B / hold B", "Distinct silhouette and accuracy benefits."),
    ("Lean", "Q / E", "Bumpers while aiming", "Collision-limited; modest accuracy cost."),
    ("Traversal", "Space", "A", "Contextual vault/mantle with stamina cost."),
    ("Fire / aim", "Mouse 1 / Mouse 2", "RT / LT", "Hitscan baseline with ballistics presentation."),
    ("Reload / interact", "R / F", "X / X", "Hold interactions for treatment or sabotage."),
    ("Grenade", "G", "RB or D-pad", "Two carried by default; clear throw preview is optional accessibility aid."),
    ("Ping / command / map", "Middle mouse / M", "D-pad / View", "Context ping, command wheel, and shared strategic map."),
], [1.18, 1.45, 1.4, 2.72])
h2("Locomotion Baseline")
table(["Posture", "Target speed", "Combat role"], [
    ("Walk", "400 cm/s", "Default deliberate movement."),
    ("Sprint", "700 cm/s", "Exposure tradeoff; drains 25 stamina/s."),
    ("Prone", "140 cm/s", "Low profile, 0.70 stationary spread multiplier, reduced AI sight signature."),
    ("Vault", "0.55 s; 12 stamina", "Fast low obstacle transition, maximum target height 100 cm."),
    ("Mantle", "0.85 s; 20 stamina", "Higher commitment, maximum target height 180 cm."),
], [1.3, 1.8, 3.95])
body("Stamina starts at 100, recovers at 20/s after a one-second delay, and primarily governs sprint and traversal. It should shape route choice and timing, not become a constant maintenance bar.")
h2("Weapon Baseline: Service Rifle")
table(["Parameter", "Current baseline", "Intent"], [
    ("Magazine / reserve", "30 / 90 rounds", "Three reload decisions before resupply pressure."),
    ("Rate of fire", "600 RPM", "Controllable cadence; bursts remain efficient."),
    ("Damage", "25 base", "Fast lethality without universal one-hit torso kills."),
    ("Spread", "0.90 hip / 0.08 ADS", "Strong incentive to shoulder weapon and stabilize."),
    ("Reload", "1.6 seconds", "Short but punishable; animation must clearly gate readiness."),
    ("Recoil", "0.55 pitch / 0.12 yaw", "Readable climb, lower in ADS, quick partial recovery."),
    ("Suppression", "250 cm radius / 0.35 amount", "Near misses alter AI decision quality and movement."),
], [1.55, 1.65, 3.85])
h2("Combat Rules")
for x in [
    "Muzzle obstruction is checked independently of the camera so cover cannot be shot through.",
    "Head, torso, arm, and leg zones change damage and injury outcomes.",
    "Movement, exhaustion, leaning, posture, and arm injury modify spread and sway.",
    "Weapon noise informs AI hearing and nearby squad alerts.",
    "Hit markers are restrained and configurable; world reaction, sound, and animation carry primary confirmation.",
    "Grenades create blast, fragmentation, suppression, and displacement; enemies can evade recognized threats.",
]:
    bullet(x)
h2("Readability Budget")
body("At combat intensity, the HUD should show only ammunition, health/bleeding/injury, teammate state, current objective, interaction prompt, ping, and short command state. Strategic alerts wait until immediate danger has passed unless they concern the active operation.")
page()

h1("4. Health, Injuries, Equipment, and Logistics")
h2("Damage and Medical Model")
body("The medical model creates recoverable tactical degradation rather than instant cascading failure. Damage reduces health; hit zones can cause arm or leg injury and bleeding; armor reduces specific damage until durability is exhausted. Field dressings stop bleeding. Medkits heal and, by default, treat limb injury after a committed three-second action.")
table(["Condition", "Effect", "Counterplay"], [
    ("Bleeding", "Ongoing health loss based on hit zone.", "Use a field dressing; reach cover first."),
    ("Arm injury", "1.65 spread multiplier and 0.85 degrees weapon sway.", "Medkit treatment; shorten engagement distance."),
    ("Leg injury", "Movement reduced to 70%.", "Medkit treatment; use squad/transport support."),
    ("Helmet impact", "Head damage scaled while durability remains.", "Avoid repeated exposure; armor is not immunity."),
    ("Body armor impact", "Torso damage scaled while durability remains.", "Repair/resupply between operations."),
    ("Incapacitation", "AI combatant stabilized rather than immediately removed.", "Creates casualty, rescue, and force-readiness consequences."),
], [1.35, 2.65, 2.75])
h2("Starting Field Load")
table(["Resource", "Baseline", "Pressure created"], [
    ("Field dressings", "3", "Enough to recover from several contacts; poor cover discipline compounds."),
    ("Medkits", "2", "Meaningful choice between health recovery and saving treatment for later."),
    ("Frag grenades", "2", "Valuable displacement tool, not routine room-clearing currency."),
    ("Rifle ammunition", "30 + 90", "Supports one operation with controlled fire; supply matters in extended fights."),
    ("Armor", "Helmet + body durability", "Absorbs mistakes but degrades across sustained contact."),
], [1.45, 1.35, 3.95])
h2("Field Logistics")
body("Supply is a physical and strategic resource. Bases generate or hold supply; connected sectors move it through routes and convoys; stations exchange it for ammunition, medical resources, armor, or operation support. The field transport bridges strategic supply and local action.")
table(["Transport property", "Baseline", "Design consequence"], [
    ("Road speed", "28 m/s; 36 m/s boost", "Fast enough to make travel useful, exposed enough to make route security matter."),
    ("Fuel", "100; 1 unit/km", "Supports broad travel without constant refueling."),
    ("Hull", "500", "Survivable asset whose loss should matter."),
    ("Cargo", "15 supply", "Deliberate logistics run rather than unlimited mobile armory."),
    ("Safe exit", "3 m/s maximum", "Prevents exploitative dismounts and communicates physical risk."),
], [1.5, 1.7, 3.55])
h2("Resource Economy Rules")
for x in [
    "Resources are spent for readiness, not used as an abstract score.",
    "Resupply should be fast at a secured base and constrained in contested territory.",
    "Abandoned transport, casualties, and unused cargo persist where feasible.",
    "The squad must never become permanently soft-locked by supply; a low-cost recovery route or fallback kit always exists.",
]:
    bullet(x)
page()

h1("5. AI, Factions, Morale, and Tactical Command")
h2("Combatant State Model")
table(["State", "Trigger", "Behavior"], [
    ("Patrol", "No threat", "Follow patrol route, observe, idle briefly, maintain believable spacing."),
    ("Alert / investigate", "Heard or uncertain stimulus", "Move toward last stimulus with increased readiness."),
    ("Combat", "Confirmed hostile", "Seek range, burst fire, suppress, reposition, claim cover, alert squad."),
    ("Evade explosive", "Recognized grenade threat", "Break cover logic and move outside danger radius."),
    ("Retreat", "Low readiness, ammunition, morale, or tactical disadvantage", "Withdraw to safer position or supply."),
    ("Search", "Lost target", "Sweep last-known area for a limited duration."),
    ("Return", "Search exhausted", "Rejoin patrol or assigned posture."),
    ("Dead/incapacitated", "Health outcome", "Stop combat contribution; persist casualty consequences."),
], [1.3, 2.2, 3.25])
h2("Perception and Engagement Baseline")
table(["Parameter", "Baseline"], [
    ("Sight / lose sight", "25 m / 30 m"),
    ("Peripheral vision", "70 degrees"),
    ("Sight memory", "2 seconds"),
    ("Hearing / memory", "35 m / 4 seconds"),
    ("Squad alert", "24 m"),
    ("Preferred engagement", "12 m; minimum 5.5 m; maximum 25 m"),
    ("Burst pattern", "2-4 rounds; 1.5-2.25 second recovery"),
], [2.2, 4.55])
h2("AI Design Requirements")
for x in [
    "AI must telegraph state changes with posture, voice, movement, and weapon readiness.",
    "Cover is claimed and released; multiple soldiers should not occupy the same anchor.",
    "Suppression reduces tactical quality without hard-stunning enemies.",
    "Squad alerts share actionable information, not perfect omniscience.",
    "Friendly and hostile soldiers use the same core combat rules where practical.",
    "Navigation failure must degrade to a safe hold/search behavior, never silent immobility.",
]:
    bullet(x)
h2("Morale and Readiness")
body("Combat readiness represents cohesion, ammunition confidence, recent casualty shock, and ability to execute. It should modify spread, burst timing, willingness to expose, and rout probability. Enemies route rather than fight to extinction when force preservation is plausible. Routed enemies count differently from killed enemies in mission scoring and future force availability.")
h2("Field Squad Command")
body("Human players form the core cooperative squad. In addition, an authorized player can lead up to three recruited AI field soldiers. Commands are intentionally small: Follow, Hold, Move-and-Hold, Board/Disembark, and Context Action. Command ownership must be visible and conflict-safe so two players cannot silently overwrite the same AI order.")
table(["Command", "Input expectation", "AI obligation"], [
    ("Follow", "One-step selection", "Maintain formation offset, catch up, avoid blocking sightline."),
    ("Hold", "One-step toggle", "Use nearby cover while preserving the ordered area."),
    ("Move-and-Hold", "Aim at location and confirm", "Path to valid point, face command yaw, report failure."),
    ("Board", "Context near transport", "Reserve seat, path, mount, and persist assignment."),
    ("Context", "Target ally/objective", "Revive, secure, sabotage, or defend when supported."),
], [1.35, 2.1, 3.2])
page()

h1("6. Missions and Operation First Light")
h2("Mission Taxonomy")
table(["Type", "Core objective", "Systemic variation"], [
    ("Attack", "Secure a hostile sector anchor.", "Approach route, reinforcements, support, contested capture."),
    ("Defense", "Hold a friendly anchor through pressure.", "Waves, breach progress, casualty checkpointing."),
    ("Raid", "Sabotage a target and exfiltrate.", "Detection state, reaction force, optional stealth."),
    ("Resupply", "Move cargo between connected locations.", "Ambush, route disruption, vehicle damage, time pressure."),
    ("Escort / rescue", "Preserve a convoy, squad, or casualty.", "Route choice, suppression, medical commitment."),
    ("Recon", "Observe or mark without triggering major contact.", "Visibility, patrol patterns, intelligence reward."),
], [1.2, 2.45, 3.0])
h2("Operation Rules")
body("Open-world operations have an approach phase followed by activation. Defense begins from a shorter base window (120 seconds); offensive operations use 240 seconds, adjusted by distance and expected travel speed, capped at 900 seconds. Once active, operation-specific rules govern capture, breach, waves, sabotage, exfiltration, casualties, and checkpoints.")
table(["Rule", "Baseline"], [
    ("Attack force", "3 enemies + 1 reinforcement wave of 1"),
    ("Defense force", "3 waves x 2 enemies, 6 seconds between waves"),
    ("Friendly support", "Up to 2 support soldiers"),
    ("Secure area", "6.5 m radius held for 8 seconds; contested progress decays"),
    ("Defense breach", "20 seconds of hostile control; recovery is twice as fast"),
    ("Raid exfiltration", "30 m from objective after sabotage"),
], [2.0, 4.75])
h2("Vertical Slice: Operation First Light")
callout("Purpose", "Prove the complete cooperative tactical grammar in a compact communications-facility mission: joining and synchronized objectives, navigation, split routes, stealth/contact escalation, authoritative keycard and door interaction, combat, supplies, persistence, respawn, regrouping, and extraction.")
h3("Required objective sequence")
table(["Order", "Objective ID", "Player-facing beat"], [
    ("1", "FindRedKeycard", "Search a readable side route and recover access credential."),
    ("2", "UnlockSecurityDoor", "Use the credential to enter the secure facility."),
    ("3", "EliminateGuard", "Resolve the principal combat threat."),
    ("4", "ReachExtraction", "Leave the site and commit mission outcome."),
], [0.65, 2.15, 3.95])
h3("Encounter layout")
for x in [
    "Two-route approach: fast/exposed service road and slower/concealed forest or maintenance line.",
    "Exterior landmarks establish the communications building, gate, security lights, and extraction direction.",
    "Patrolling guard creates an observation window rather than a fixed shooting gallery.",
    "Keycard and door demonstrate persistent world-state contracts.",
    "Ammo and medical supplies appear before or immediately after the highest-risk contact.",
    "Extraction requires the principal threat objective so the player cannot bypass the mission grammar accidentally.",
]:
    bullet(x)
h3("Slice success criteria")
table(["Discipline", "Exit criterion"], [
    ("Design", "A first-time squad can infer both routes, assign roles, understand the access problem, and finish in 15-25 minutes."),
    ("Combat", "Guard transitions through perception states, uses cover/repositioning, and produces readable damage feedback."),
    ("Networking", "Remote players receive combat, objective, door, keycard, supply, and operation state without duplication or divergence."),
    ("Persistence", "Authoritative keycard, door, objectives, inventory, operation, and field-respawn state restore consistently."),
    ("Art", "Cold wet dawn, concrete/steel facility, forest silhouette, orange security accents, no route-obscuring clutter."),
    ("Audio", "Distant war bed, weather, facility machinery, positional gunfire, and interaction confirmations are distinct."),
    ("Performance", "Target frame-time achieved on reference PC with scalable fog, shadows, reflections, and foliage."),
], [1.25, 5.5])
page()

h1("7. Persistent War, Sectors, Routes, and Strategic Map")
h2("Strategic Entities")
table(["Entity", "State", "Player meaning"], [
    ("Sector", "Owner, pressure, strength, supply, priority, anchor", "Territory that creates missions and route access."),
    ("Route", "Endpoints, connection, control, risk", "Travel and supply link that can be disrupted."),
    ("Supply base", "Stock, production/recovery capability", "Source of readiness and reinforcement."),
    ("Convoy", "Source, destination, cargo, survival", "Physical transfer of strategic supply."),
    ("Operation", "Type, phase, forces, progress, casualties", "Current tactical expression of the war."),
    ("War event", "Turn, location, result, consequence", "Audit trail that explains why the map changed."),
], [1.25, 2.45, 3.05])
h2("Priority Generation")
body("The war system evaluates sectors using ownership, adjacency, pressure, supply connectivity, strategic weight, and recent events. It proposes defend, attack, raid, resupply, or militia actions. Priority scoring must remain inspectable in debug and explainable in UI with two or three human-readable reasons.")
callout("Map principle", "The strategic map is a shared decision surface, not wallpaper. Every squad member must see the same revision, and every icon must answer: what is happening, why it matters, how long we have, what it costs, and what success changes.")
h2("Strategic Map Information Architecture")
table(["Layer", "Default", "On selection"], [
    ("Territory", "Sector ownership and contested state", "Strength, pressure, adjacent sectors, recent change"),
    ("Supply", "Connected/disconnected route emphasis", "Stock, source, convoy, blockage, resupply options"),
    ("Operations", "Primary and secondary priority markers", "Type, approach time, forces, likely consequences"),
    ("Player assets", "All players, AI squad, transport", "Status, separation, fuel, cargo, hull, assignments, route estimate"),
    ("History", "Last major change only", "Short event log with causal language"),
], [1.25, 2.55, 2.95])
h2("Persistence Contract")
for x in [
    "Stable persistence IDs are mandatory for world actors and field assets.",
    "Objective IDs and reflected names are content/save contracts.",
    "Campaign data includes per-player state, shared objective/keycard state, persistent actors, operation phase/progress, war state, AI squad assignment, and transport state.",
    "Authoritative save writes occur on explicit checkpoints, operation phase changes, material casualties, travel commits, and controlled host shutdown where safe.",
    "Version migrations favor conservative defaults and never silently reinterpret an ID.",
]:
    bullet(x)
h2("Campaign Balance Guardrails")
table(["Guardrail", "Target"], [
    ("Critical priorities simultaneously visible", "1 primary + up to 2 secondary"),
    ("Unanswered sector loss cadence", "No faster than one meaningful loss per 2-3 completed operations in standard difficulty"),
    ("Recovery from isolation", "At least one achievable relief/resupply option"),
    ("War explanation", "Every ownership change produces a player-readable event"),
    ("Simulation determinism", "Seeded or logged enough to reproduce QA cases"),
], [2.7, 4.05])
page()

h1("8. World, Level, Art, and Audio Direction")
h2("World Structure")
body("The game uses a connected regional world centered on sectors, routes, forward bases, industrial sites, settlements, forests, passes, and communications infrastructure. Traversal space must earn its existence through sightline decisions, route risk, logistics, ambient war activity, or environmental storytelling. Empty distance without decision value is compressed.")
h2("Level Design Grammar")
table(["Scale", "Required ingredients"], [
    ("Regional", "Recognizable horizon landmarks, route loops, supply connections, controlled streaming, plausible travel times."),
    ("Sector", "Anchor, entry vectors, defensible geometry, flanking path, logistics relationship, civilian or industrial context."),
    ("Encounter", "Cover chain, exposure crossing, vertical option, retreat route, readable objective, AI navigation redundancy."),
    ("Interaction", "Clear affordance, safe approach space, feedback, persistence ID when state survives reload."),
], [1.25, 5.5])
h2("Art Direction")
body("The visual target is a cold, wet dawn in a fictional Eastern European mountain region. Architecture is functional: weathered concrete, galvanized or painted steel, worn asphalt, service roads, drainage, cables, antennas, maintenance access, and practical defensive additions. Technology is incremental and rugged rather than futuristic.")
table(["Category", "Use", "Avoid"], [
    ("Palette", "Charcoal, gunmetal, muted olive, steel blue, concrete gray, desaturated forest", "Oversaturated green, neon color coding"),
    ("Light", "Overcast dawn, subdued reflection, sparse burnt-orange security lamps", "Excess bloom, crushed blacks, omnipresent emissives"),
    ("Material", "Wet concrete, steel, mud, moss, timber, worn asphalt", "Glossy sci-fi surfaces"),
    ("Composition", "Landmarks, readable paths, tactical silhouettes, protected sightlines", "Clutter that hides routes, pickups, prompts, or targets"),
    ("Identity", "Fictional symbols and neutral practical markings", "Real flags, insignia, political slogans"),
], [1.15, 3.05, 2.55])
h2("Audio Direction")
table(["Layer", "Design goal"], [
    ("Weapons", "Sharp mechanical identity, indoor/outdoor tails, suppressed near-miss language, no cinematic over-compression."),
    ("Movement", "Surface, posture, equipment, and speed communicate risk to player and AI."),
    ("Weather/world", "Wind, rain, trees, distant machinery, generators, cable hum, structural resonance."),
    ("War bed", "Distant artillery, aircraft, and intermittent small arms imply a wider front without constant spectacle."),
    ("AI voice", "Short fictional-language or neutral barks for alert, contact, reload, grenade, casualty, retreat, and search."),
    ("UI", "Quiet confirmation; strategic warnings are distinct from immediate combat alarms."),
], [1.4, 5.35])
h2("Music")
body("Music is sparse. Textural low strings, processed percussion, and restrained electronic pulses support approach and consequence. Combat music should not reveal enemies before the simulation does. Silence and environmental sound are primary tension tools.")
h2("Performance-Aware Art Rules")
for x in [
    "Modular kits, trim sheets, shared material instances, decals, HISM, and disciplined LODs.",
    "Conservative shadow-casting lights; orange accents identify function, not decoration.",
    "Scalable fog, reflection, foliage, and shadow tiers with a stable readability floor.",
    "World Partition and navigation invokers must preserve arrival and combat reliability.",
]:
    bullet(x)
page()

h1("9. UI/UX, Onboarding, Accessibility, and Settings")
h2("HUD Hierarchy")
table(["Priority", "Element", "Behavior"], [
    ("Immediate", "Crosshair/weapon state, ammo, injury/bleeding, interaction", "Visible during relevant action; strong but compact."),
    ("Tactical", "Teammate state, ping, current objective, squad command, grenade threat, hit direction", "Appears contextually and fades."),
    ("Operational", "Capture/breach/wave/exfiltration progress", "Visible only during active operation."),
    ("Strategic", "Priority, sector change, supply event", "Deferred until safe unless time-critical."),
], [1.0, 2.45, 3.45])
h2("Screen Set")
for x in [
    "Main menu: Host Campaign, Join Session, Server Browser/Invite, Operation First Light, Settings, Credits, Quit.",
    "In-session menu: Resume, squad/connection state, objective summary, controls, settings, leave session; it does not pause the server.",
    "Strategic map: sector, route, operation, supply, and field-asset layers.",
    "After-action report: objective result, casualties, ammunition/medical usage, territory and supply consequences.",
    "Settings: display, graphics, audio, controls, gameplay, accessibility.",
]:
    bullet(x)
h2("Onboarding")
table(["Stage", "Teaches", "Method"], [
    ("Safe movement", "Look, move, sprint, crouch, prone, lean, interact", "Short context prompts with immediate practice."),
    ("First Light approach", "Observation, alternate route, AI perception", "Environmental framing; no forced stealth."),
    ("Access problem", "Keycard, locked door, persistence", "Single clear credential loop."),
    ("First contact", "ADS, recoil, suppression, cover, medical state", "One principal guard with readable patrol."),
    ("Extraction", "Objective order, checkpoint, consequence", "Explicit commit and after-action summary."),
    ("World transition", "Shared map, priorities, supply, travel", "Guided first squad vote/host commitment after slice."),
], [1.4, 2.5, 3.0])
h2("Accessibility Requirements")
for x in [
    "Full key and controller remapping; independent horizontal/vertical and ADS sensitivity.",
    "Subtitles with speaker labels, scalable size, directional indicators, and background opacity.",
    "Color-independent ownership and injury cues using icons, patterns, text, and audio.",
    "Scalable HUD and text; safe-area controls; high-contrast interaction mode.",
    "Toggle/hold options for aim, sprint, crouch, prone, lean, and interactions.",
    "Camera shake, recoil animation, head bob, hit flash, motion blur, depth of field, and chromatic aberration controls.",
    "Difficulty assists: aim friction, enemy perception, incoming damage, treatment duration, strategic pressure, grenade trajectory preview.",
    "Menus never pause a network session; important briefings provide a ready check, safe staging area, or host-controlled planning hold.",
]:
    bullet(x)
h2("UX Acceptance Tests")
table(["Test", "Pass condition"], [
    ("Objective comprehension", "80% of first-time squads agree on the next action within 15 seconds."),
    ("Damage comprehension", "Players distinguish bleeding, limb injury, armor impact, and low health."),
    ("Map decision", "Players can identify the primary crisis, consequence, and host commitment state within 20 seconds."),
    ("Prompt overload", "No more than one interaction prompt and one strategic notification compete."),
    ("Color vision", "All state changes remain interpretable in monochrome capture."),
], [2.0, 4.75])
page()

h1("10. Progression, Difficulty, Economy, and Balance")
h2("Progression Philosophy")
body("Progression expands tactical options and campaign resilience rather than inflating damage numbers. The player becomes more capable through information, field relationships, recovered assets, improved support, and mastery of systems.")
table(["Track", "Unlock examples", "Prohibited outcome"], [
    ("Equipment access", "Optics, utility tools, armor repair, alternate ammunition roles", "Straight rarity ladder invalidating baseline rifle"),
    ("Field capability", "Larger support availability, faster casualty recovery, better transport service", "Automated victory without player action"),
    ("Intelligence", "Better route risk, force estimate, patrol and supply visibility", "Perfect information"),
    ("Regional trust", "Militia mobilization, local caches, reconnaissance reports", "Moral choices reduced to a currency meter"),
    ("Player mastery", "No unlock required: lean, posture, suppression, route planning", "Tutorial-gated basic competence"),
], [1.45, 2.95, 2.35])
h2("Difficulty Model")
table(["Axis", "Recruit", "Operator", "Veteran"], [
    ("Incoming damage", "0.75x", "1.0x", "1.15x"),
    ("Enemy perception", "Reduced acquisition", "Baseline", "Faster, not omniscient"),
    ("Enemy coordination", "Longer delays", "Baseline", "Shorter alerts/reposition"),
    ("Medical pressure", "More supplies, faster treatment", "Baseline", "Scarcer field recovery"),
    ("Strategic pressure", "Slow losses, generous approach", "Baseline", "Faster pressure, same information quality"),
    ("Checkpointing", "Frequent", "Phase-based", "Phase-based; no punitive save deletion"),
], [1.45, 1.65, 1.65, 2.0])
body("Difficulty never removes key feedback, makes AI cheat, or changes persistence contracts. Custom difficulty exposes each axis.")
h2("Scoring and After-Action Evaluation")
table(["Dimension", "Examples"], [
    ("Mission result", "Success, partial success, retreat, failure"),
    ("Force preservation", "Friendly casualties, stabilized casualties, transport survival"),
    ("Enemy outcome", "Killed, incapacitated, routed, bypassed"),
    ("Resource efficiency", "Ammunition, medical supplies, armor, cargo, fuel"),
    ("Operational effect", "Sector control, supply connectivity, intelligence, pressure"),
    ("Conduct", "Optional civilian/infrastructure preservation when relevant"),
], [1.8, 4.95])
h2("Balance Targets")
for x in [
    "Standard rifle torso time-to-disable: approximately 3-5 accurate hits depending on armor and zone.",
    "Typical contact: 2-6 active hostiles, with reinforcements staged rather than all visible at once.",
    "First Light completion: 15-25 minutes first play, 8-15 minutes repeat play.",
    "Standard operation: 20-45 minutes including approach; travel should be less than one-third unless events intervene.",
    "Medical resource use: an average successful operation consumes 0-2 dressings and 0-1 medkits.",
    "Campaign losses should emerge from accumulated poor choices, not one opaque simulation roll.",
]:
    bullet(x)
h2("Ethical and Content Boundaries")
body("Broken Horizon avoids real-world faction mapping, celebratory violence, and dehumanizing enemy presentation. Combatants communicate fear, suppression, casualty, and retreat. The design may portray the cost of conflict, but does not reward cruelty or use civilian harm as disposable spectacle.")
page()

h1("11. Technical Architecture, Content Pipeline, Analytics, and QA")
h2("Engine and Runtime")
table(["Area", "Current direction"], [
    ("Engine", "Unreal Engine 5.8; C++ runtime module with Blueprint-facing APIs."),
    ("Networking", "Server-authoritative cooperative play; replicated GameState snapshots, components, actors, RPCs, multicast presentation, listen and dedicated server paths."),
    ("Input", "Enhanced Input."),
    ("AI", "AI Perception, navigation, StateTree/GameStateTree dependencies, source-backed tactical controller."),
    ("World", "World Partition/open regional map, dynamic navigation, navigation invokers."),
    ("UI", "Native BH widget classes plus UMG Blueprint assets."),
    ("FX", "Niagara-supported presentation; scalable effects."),
    ("Save", "BHSaveGame + BHSaveSubsystem and separate user settings persistence."),
    ("Editor", "Dedicated editor module and idempotent Python automation/validators."),
], [1.55, 5.2])
h2("System Ownership")
table(["System", "Primary classes / data"], [
    ("Players", "Replicated BHCharacter, BHWeaponComponent, BHRifle, BHHealthComponent, BHInjuryComponent"),
    ("AI", "BHEnemySoldier, BHEnemyAIController, patrol and cover actors"),
    ("Mission", "BHMissionData, BHObjectiveComponent, operation/defense directors, extraction"),
    ("War", "BHWarSubsystem, BHWarTypes, BHWarOperationRules, BHWarGameState"),
    ("World logistics", "BHSectorAnchor, BHWorldRoute, BHSupplyBase, BHSupplyConvoyTarget, BHFieldTransport"),
    ("Persistence", "Authority-owned BHSaveGame/BHSaveSubsystem, persistent actor IDs, per-player state, and replicated operation snapshots"),
    ("UI", "BHWarMapWidget, objective, combat, ammo, interaction, menu, death, and mission-complete widgets"),
], [1.55, 5.2])
h2("Content Authoring Contracts")
for x in [
    "Runtime code remains in Source/BrokenHorizon; editor-only automation remains in Source/BrokenHorizonEditor or Content/Python.",
    "Never hand-edit binary .uasset or .umap files.",
    "Reflected property names, asset paths, objective IDs, collision profiles, and persistence IDs require consumer search before change.",
    "Persistent world actors receive unique stable IDs.",
    "Editor automation is idempotent, tagged, narrowly scoped, and refuses destructive overwrite by default.",
    "Generated First Light content uses the BH_Auto_FirstLight tag.",
]:
    bullet(x)
h2("Telemetry / Playtest Instrumentation")
body("For internal builds, log anonymized local playtest events to a developer-readable file. Online collection requires explicit product approval and consent.")
table(["Event", "Fields"], [
    ("Operation lifecycle", "Type, sector, phase times, result, retries, route"),
    ("Combat", "Contact count, engagement range, shots, hits by zone, suppression, damage source"),
    ("Medical", "Bleed onset, treatment delay, resources spent, death cause"),
    ("AI", "State time, perception source, cover failures, navigation failures, rout"),
    ("War", "Priority chosen/ignored, sector changes, supply path changes"),
    ("UX", "Objective dwell, map dwell, settings changed, prompt dismissal"),
], [1.65, 5.1])
h2("Validation Matrix")
table(["Change area", "Automated evidence", "Manual evidence"], [
    ("Gameplay C++", "Editor build, automation tests, targeted runtime smoke", "Hands-on feel and regression play"),
    ("Mission/AI/combat/save", "Build + tests + smoke + First Light launch/log review", "Navigation, Blueprint wiring, encounter quality"),
    ("UI/visual/audio", "Load and asset validation", "Resolution sweep, input sweep, visibility, mix"),
    ("Startup/cook/package", "Packaged build launch and log review", "Fresh-machine install and full flow"),
    ("Persistence", "Save/load automation across state permutations", "Long-session campaign and migration soak"),
    ("Networking", "Dedicated/listen server launch, multi-client functional tests, replication assertions", "Latency/loss simulation, join-in-progress, reconnect, host travel, field respawn"),
], [1.55, 2.6, 2.6])
page()

h1("12. Production Scope, Roadmap, Risks, and Acceptance Criteria")
h2("Recommended Product Scope")
table(["Tier", "Included"], [
    ("Vertical slice", "Networked Operation First Light, synchronized combat/medical/objective loop, join-in-progress, field respawn, authority-owned persistence, polished presentation."),
    ("Minimum viable campaign", "Shared connected region, 6-8 sectors, 3 core operation types, human squad plus AI support, transport/logistics, strategic map, 4-6 hour arc."),
    ("1.0 target", "Stable cooperative campaign across listen/dedicated server paths, 10-14 sectors, 5-6 operation types, expanded equipment and AI roles, 10-15 hour arc."),
    ("Post-1.0 option", "Expanded squad scale, server administration, additional campaign roles, and platform cross-play after the core network experience is proven."),
], [1.65, 5.1])
h2("Roadmap Gates")
table(["Gate", "Exit criteria"], [
    ("G0 - Foundations", "Build/test loop reliable; authority and persistence contracts documented; replicated player, weapon, injury, AI, vehicle, and war state functional."),
    ("G1 - First Light alpha", "Two or more clients complete the mission end-to-end with synchronized objectives, combat, field respawn, and save state."),
    ("G2 - First Light vertical slice", "Join/reconnect, listen/dedicated paths, target art/audio, tuned combat, accessibility, network/performance, and packaging pass."),
    ("G3 - Campaign prototype", "Connected clients share three sectors, attack/defense/raid, supply route, transport, server travel, and persistent consequences."),
    ("G4 - Content production", "Authoring templates, validators, encounter budget, save migration, performance budgets locked."),
    ("G5 - Alpha", "All systems/content playable in supported multiplayer topology; no blocking desync or save loss; complete campaign flow."),
    ("G6 - Beta", "Content locked; balance, accessibility, performance, compatibility, localization, and packaging focus."),
    ("G7 - Release candidate", "Zero critical defects, approved known-issue list, full multiplayer regression, migration, disconnect, recovery, and soak tests."),
], [1.65, 5.1])
h2("Top Risks and Mitigations")
table(["Risk", "Impact", "Mitigation"], [
    ("Persistent-war complexity", "Opaque outcomes and difficult balancing", "Explainable scores, deterministic logs, limited simultaneous priorities."),
    ("Open-world AI navigation", "Stalls, poor cover, broken operations", "Invoker coverage validation, fallback behaviors, encounter nav budgets."),
    ("Save-state breadth", "Corruption or incompatible contracts", "Stable IDs, versioning, migration tests, staged checkpoints, backups."),
    ("Scope inflation", "Many shallow systems", "First Light quality gate; operation types reuse shared rules."),
    ("Readability vs realism", "Dark scenes and ambiguous damage", "Minimum visibility floor, layered feedback, accessibility controls."),
    ("Vehicle/logistics friction", "Travel becomes chore", "Fast road speed, route events, recovery options, concise service interactions."),
    ("Network state divergence", "Clients see different war, objective, inventory, or transport state", "Server authority, revisioned snapshots, idempotent actions, multi-client functional tests."),
    ("Host/session loss", "Campaign interruption or lost progress", "Controlled autosaves, reconnect policy, explicit host-loss messaging; host migration is a product gate."),
], [1.65, 2.2, 2.9])
h2("Definition of Done: Player Experience")
for x in [
    "Every connected player sees the same current tactical objective, operation phase, and strategic reason for it.",
    "Movement, posture, recoil, suppression, injuries, and treatment create distinct, learnable tradeoffs.",
    "Enemies perceive, communicate, use cover, reposition, search, evade grenades, and retreat without obvious cheating.",
    "At least two viable tactical approaches exist for principal encounters.",
    "Operation outcomes persist authoritatively and visibly alter territory, supply, forces, or future priorities for all clients.",
    "The game remains comprehensible with color removed and core camera effects disabled.",
    "A complete multiplayer campaign can be finished without desync, duplicated state, save loss, unrecoverable supply lock, or mandatory debug intervention.",
]:
    bullet(x)
h2("Open Product Decisions")
table(["Decision", "Recommended default", "Decision gate"], [
    ("Player cap", "Squad-sized; lock the exact cap through First Light replication and encounter tests", "Before campaign content budgets"),
    ("Session discovery", "Direct invite plus server browser; matchmaking only if audience testing supports it", "Before online services integration"),
    ("Host migration", "Not assumed; evaluate against campaign ownership and platform services", "Before persistence architecture lock"),
    ("Player identity", "Role/loadout identity with light visual customization", "Before narrative/content production"),
    ("Console", "PC first; maintain controller/performance compatibility", "After vertical-slice profiling"),
    ("Mission restart policy", "Checkpoint retry plus explicit persistent-outcome option", "Campaign prototype playtest"),
    ("Campaign length", "10-15 hours for 1.0", "After three-sector pacing data"),
    ("Voice language", "Fictionalized/neutral with localized subtitles", "Before audio casting"),
], [1.65, 2.9, 2.2])
callout("Next production action", "Treat networked Operation First Light as the quality bar. Do not expand regional content until multiple clients can join, fight, interact, respawn, reconnect, travel, save, and extract with synchronized state, strong readability, stable performance, and target presentation.", fill="FCEFE7")

h1("Appendix A. Current Baseline Tuning Reference")
table(["System", "Selected baseline"], [
    ("Player movement", "Walk 400; sprint 700; prone 140 cm/s; stamina 100, drain 25/s, recovery 20/s"),
    ("Traversal", "Vault max 100 cm / 0.55 s / 12 stamina; mantle max 180 cm / 0.85 s / 20 stamina"),
    ("Rifle", "30+90, 600 RPM, 25 damage, 1.6 s reload, 50,000 cm range"),
    ("Medical", "3 dressings, 2 medkits, 45 heal, 3.0 s treatment"),
    ("Posture/injury", "Prone spread 0.70; arm injury spread 1.65; leg movement 0.70"),
    ("AI perception", "Sight 25 m, lose 30 m, hearing 35 m, squad alert 24 m"),
    ("AI fire", "2-4 shot bursts, 1.5-2.25 s recovery, desired range 12 m"),
    ("Operations", "Attack 3+1; defense 3x2; secure 8 s; breach 20 s"),
    ("Transport", "28 m/s, boost 36 m/s, 100 fuel, 500 hull, 15 cargo"),
], [1.55, 5.2])
h2("Appendix B. Canonical First Light Contracts")
table(["Contract", "Value"], [
    ("Map", "/Game/BrokenHorizon/Maps/L_FirstLight_Graybox"),
    ("Generated actor tag", "BH_Auto_FirstLight"),
    ("Keycard ID / persistence", "RedKeycard / FirstLightRedKeycard"),
    ("Door persistence", "FirstLightSecurityDoor"),
    ("Objective order", "FindRedKeycard > UnlockSecurityDoor > EliminateGuard > ReachExtraction"),
    ("Supplies", "FirstLightAmmoSupply / FirstLightMedicalSupply"),
], [2.15, 4.6])
h2("Appendix C. Source Basis")
body("This GDD was synthesized from the current project configuration, source headers and implementations, project automation documentation, and the Operation First Light art-direction package. Numeric values are implementation baselines, not immutable promises. Proposed scope and content breadth require production approval and playtest validation.")

doc.core_properties.title = "Broken Horizon - Game Design Document"
doc.core_properties.subject = "Implementation-informed production game design document"
doc.core_properties.author = "Broken Horizon Development Team"
doc.core_properties.keywords = "Broken Horizon, GDD, tactical FPS, persistent war, Unreal Engine"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
